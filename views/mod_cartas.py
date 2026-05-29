import streamlit as st
import pandas as pd
import zipfile
import os
from datetime import datetime
from docx import Document
from io import BytesIO
from modulos import database as db


# ─── GERAÇÃO DE WORD ──────────────────────────────────────────────────────────

def gerar_word_memoria(dados):
    try:
        diretorio_raiz = os.getcwd()
        template_path = os.path.join(diretorio_raiz, "carta_preenchida.docx")
        if not os.path.exists(template_path):
            template_path = "carta_preenchida.docx"

        doc = Document(template_path)

        for p in doc.paragraphs:
            for k, v in dados.items():
                if f"{{{{{k}}}}}" in p.text:
                    p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in dados.items():
                            if f"{{{{{k}}}}}" in p.text:
                                p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"⚠️ Erro no Template Word: {e}")
        return None


# ─── HELPERS DE DADOS ────────────────────────────────────────────────────────

def obter_base_colaboradores(fire):
    docs = fire.collection("colaboradores_base").stream()
    return {doc.id: doc.to_dict().get("cpf") for doc in docs}


def salvar_novo_colaborador(fire, nome, cpf):
    fire.collection("colaboradores_base").document(nome.upper().strip()).set(
        {"cpf": str(cpf).strip()}
    )


def _dados_word(c):
    return {
        "NOME_COLAB":    c["NOME"],
        "CPF":           c["CPF"],
        "CODIGO_CLIENTE": c.get("COD_CLI", ""),
        "VALOR_DEBITO":  f"R$ {c['VALOR']:,.2f}",
        "LOJA_ORIGEM":   c.get("LOJA", ""),
        "DATA_COMPRA":   c.get("DATA", ""),
        "DESC_DEBITO":   c.get("MOTIVO", ""),
        "DATA_LOCAL":    f"São Paulo, {datetime.now().strftime('%d/%m/%Y')}",
    }


# ─── GERAÇÃO DE ZIP (todos os .docx de um lote) ───────────────────────────────

def gerar_zip_lote(cartas_lote):
    """Recebe lista de dicts de cartas e retorna bytes de um ZIP com todos os .docx."""
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for c in cartas_lote:
            w_bytes = gerar_word_memoria(_dados_word(c))
            if w_bytes:
                nome_arquivo = f"Carta_{c['NOME'].replace(' ', '_')}_{c.get('id','')}.docx"
                zf.writestr(nome_arquivo, w_bytes)
    zip_buffer.seek(0)
    return zip_buffer.getvalue()


# ─── GERAÇÃO DE EXCEL ────────────────────────────────────────────────────────

def gerar_excel_lote(cartas_lote):
    """Recebe lista de dicts e retorna bytes de um .xlsx."""
    colunas = ["NOME", "CPF", "COD_CLI", "VALOR", "LOJA", "DATA", "MOTIVO", "status", "id_lote"]
    rows = []
    for c in cartas_lote:
        rows.append({col: c.get(col, "") for col in colunas})

    df = pd.DataFrame(rows, columns=colunas)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Cartas")
        ws = writer.sheets["Cartas"]
        # Ajusta largura das colunas automaticamente
        for i, col in enumerate(df.columns):
            largura = max(df[col].astype(str).map(len).max(), len(col)) + 2
            ws.set_column(i, i, largura)
    buf.seek(0)
    return buf.getvalue()


# ─── INTERFACE PRINCIPAL ─────────────────────────────────────────────────────

def exibir(user_role):
    st.markdown("""
        <style>
        .metric-card {
            background-color: #F8F9FA; padding: 15px; border-radius: 10px;
            border-left: 5px solid #D4AF37; text-align: center;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        .card-carta {
            background-color: #F8F9FA; padding: 20px; border-radius: 12px;
            border-top: 5px solid #D4AF37; box-shadow: 0 4px 6px rgba(0,0,0,0.05);
            margin-bottom: 20px; min-height: 180px;
        }
        .loja-header {
            background-color: #495057; color: #D4AF37; padding: 8px 15px;
            border-radius: 8px; margin: 25px 0 15px 0; font-weight: bold;
        }
        .label-card { color: #6c757d; font-size: 0.85rem; margin-bottom: 2px; }
        .info-card  { font-weight: 600; color: #212529; margin-bottom: 8px; }
        </style>
    """, unsafe_allow_html=True)

    st.title("📑 Gestão de Cartas de Débito")

    fire = db.inicializar_db()
    if not fire:
        st.error("Sem conexão com o banco de dados.")
        return

    # ── Carrega dados frescos a cada render ──────────────────────────────────
    def obter_cartas():
        return [doc.to_dict() for doc in fire.collection("cartas_rh").stream()]

    cartas       = obter_cartas()
    dict_colab   = obter_base_colaboradores(fire)
    lista_nomes  = sorted(list(dict_colab.keys()))

    tabs = st.tabs(["🆕 Nova Carta", "📋 Painel", "📦 Fechamento", "✅ Histórico", "⚙️ Config"])

    # ── ABA 0: NOVA CARTA ────────────────────────────────────────────────────
    with tabs[0]:
        st.subheader("Informações do Lançamento")
        escolha_nome = st.selectbox(
            "Busque ou selecione o Colaborador:", ["+ CADASTRAR NOVO"] + lista_nomes
        )

        with st.form("f_premium", clear_on_submit=True):
            c1, c2, c3 = st.columns(3)
            if escolha_nome == "+ CADASTRAR NOVO":
                nome   = c1.text_input("Nome do Colaborador").upper().strip()
                cpf    = c2.text_input("CPF")
            else:
                nome = escolha_nome
                cpf  = c2.text_input("CPF", value=dict_colab.get(escolha_nome, ""))

            cod_cli = c3.text_input("Código do Cliente")
            v1, v2, v3 = st.columns(3)
            valor  = v1.number_input("Valor R$", min_value=0.0, step=0.01)
            loja   = v2.text_input("Loja Origem").upper()
            data_c = v3.date_input("Data da Ocorrência")
            motivo = st.text_area("Motivo Detalhado").upper()

            if st.form_submit_button("✨ Gerar e Registrar", type="primary", use_container_width=True):
                if nome and cpf and cod_cli:
                    if escolha_nome == "+ CADASTRAR NOVO":
                        salvar_novo_colaborador(fire, nome, cpf)

                    id_carta  = datetime.now().strftime("%Y%m%d%H%M%S")
                    dados_fb  = {
                        "id": id_carta, "NOME": nome, "CPF": cpf,
                        "COD_CLI": cod_cli, "VALOR": valor, "LOJA": loja,
                        "DATA": data_c.strftime("%d/%m/%Y"), "MOTIVO": motivo,
                        "status": "Aguardando Assinatura", "anexo_bin": None,
                        "data_criacao": datetime.now().strftime("%d/%m/%Y %H:%M"),
                    }
                    fire.collection("cartas_rh").document(id_carta).set(dados_fb)
                    st.success("✅ Carta registrada com sucesso!")
                    st.rerun()
                else:
                    st.error("Preencha Nome, CPF e Código do Cliente.")

    # ── ABA 1: PAINEL ────────────────────────────────────────────────────────
    with tabs[1]:
        lista_painel = [c for c in cartas if c.get("status") == "Aguardando Assinatura"]

        if lista_painel:
            total_valor = sum(c["VALOR"] for c in lista_painel)
            m1, m2 = st.columns(2)
            m1.markdown(
                f'<div class="metric-card"><div class="label-card">Pendentes</div>'
                f'<div style="font-size:1.5rem;font-weight:bold;">{len(lista_painel)}</div></div>',
                unsafe_allow_html=True,
            )
            m2.markdown(
                f'<div class="metric-card"><div class="label-card">Valor Total</div>'
                f'<div style="font-size:1.5rem;font-weight:bold;color:#D4AF37;">'
                f'R$ {total_valor:,.2f}</div></div>',
                unsafe_allow_html=True,
            )
            st.write("---")

            busca_p = st.text_input("🔍 Buscar (Nome / Código)")
            if busca_p:
                lista_painel = [
                    c for c in lista_painel
                    if busca_p.upper() in c["NOME"] or busca_p in str(c.get("COD_CLI", ""))
                ]

            df_p = pd.DataFrame(lista_painel).sort_values(by="LOJA")

            for loja_n, group in df_p.groupby("LOJA"):
                st.markdown(
                    f'<div class="loja-header">📍 {loja_n} ({len(group)} itens)</div>',
                    unsafe_allow_html=True,
                )
                cols = st.columns(3)
                for idx, (_, c) in enumerate(group.iterrows()):
                    with cols[idx % 3]:
                        st.markdown(
                            f'<div class="card-carta">'
                            f'<div class="label-card">COLABORADOR</div><div class="info-card">{c["NOME"]}</div>'
                            f'<div class="label-card">CÓD. CLIENTE</div><div class="info-card">{c.get("COD_CLI","—")}</div>'
                            f'<div class="label-card">VALOR</div>'
                            f'<div style="font-size:1.2rem;font-weight:bold;color:#D4AF37;">R$ {c["VALOR"]:,.2f}</div>'
                            f'<div class="label-card">DATA: {c["DATA"]}</div>'
                            f'<div class="label-card" style="margin-top:8px;">MOTIVO</div>'
                            f'<div style="font-size:0.85rem;color:#495057;">{c.get("MOTIVO","—")}</div>'
                            f"</div>",
                            unsafe_allow_html=True,
                        )

                        w_bytes = gerar_word_memoria(_dados_word(c))
                        btn_c1, btn_c2 = st.columns(2)

                        if w_bytes:
                            btn_c1.download_button(
                                "📂 Baixar",
                                w_bytes,
                                file_name=f"Carta_{c['NOME']}.docx",
                                key=f"w_{c['id']}",
                                use_container_width=True,
                            )

                        if user_role in ["ADM", "GERENTE"]:
                            if btn_c2.button("🗑️", key=f"del_{c['id']}", use_container_width=True):
                                fire.collection("cartas_rh").document(c["id"]).delete()
                                st.rerun()

                        up = st.file_uploader(
                            "Upload Assinada", key=f"up_{c['id']}", label_visibility="collapsed"
                        )
                        if up:
                            fire.collection("cartas_rh").document(c["id"]).update({
                                "status":       "CARTA RECEBIDA",
                                "anexo_bin":    up.getvalue(),
                                "nome_arquivo": up.name,
                            })
                            st.success("✅ Recebida!")
                            st.rerun()
        else:
            st.info("Nenhuma carta aguardando assinatura.")

    # ── ABA 2: FECHAMENTO DE LOTE ────────────────────────────────────────────
    with tabs[2]:
        prontas = [c for c in cartas if c.get("status") == "CARTA RECEBIDA"]

        if not prontas:
            st.info("Nenhuma carta assinada pronta para fechamento.")
        else:
            st.subheader(f"📦 Lote pronto: {len(prontas)} itens")

            df_prontas = pd.DataFrame(prontas)[["NOME", "VALOR", "LOJA", "COD_CLI", "DATA"]]
            st.dataframe(df_prontas, use_container_width=True)

            total_lote = sum(c["VALOR"] for c in prontas)
            st.metric("Valor Total do Lote", f"R$ {total_lote:,.2f}")
            st.divider()

            # ── Downloads ANTES de fechar (pré-visualização) ──────────────
            col_zip, col_xls = st.columns(2)

            zip_bytes = gerar_zip_lote(prontas)
            col_zip.download_button(
                "📥 Baixar ZIP (todos os .docx)",
                data=zip_bytes,
                file_name=f"Lote_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
                mime="application/zip",
                use_container_width=True,
            )

            excel_bytes = gerar_excel_lote(prontas)
            col_xls.download_button(
                "📊 Baixar Excel do Lote",
                data=excel_bytes,
                file_name=f"Lote_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

            st.divider()

            if st.button(
                "🚀 FINALIZAR LOTE E ENVIAR AO HISTÓRICO",
                type="primary",
                use_container_width=True,
            ):
                id_lote    = datetime.now().strftime("%Y%m%d_%H%M")
                ids_cartas = [c["id"] for c in prontas]

                fire.collection("lotes_rh").document(id_lote).set({
                    "id":        id_lote,
                    "data":      datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "total":     len(prontas),
                    "valor_total": total_lote,
                    "ids_cartas": ids_cartas,
                })

                for id_c in ids_cartas:
                    fire.collection("cartas_rh").document(id_c).update({
                        "status":   "LOTE_FECHADO",
                        "id_lote":  id_lote,
                    })

                st.success(f"✅ Lote {id_lote} finalizado! {len(prontas)} cartas arquivadas.")
                st.rerun()

    # ── ABA 3: HISTÓRICO ────────────────────────────────────────────────────
    with tabs[3]:
        st.subheader("✅ Histórico de Lotes Fechados")

        lotes_docs = list(fire.collection("lotes_rh").stream())

        if not lotes_docs:
            st.info("Nenhum lote fechado ainda.")
        else:
            lotes = sorted(
                [doc.to_dict() for doc in lotes_docs],
                key=lambda x: x.get("id", ""),
                reverse=True,
            )

            for lote in lotes:
                with st.expander(
                    f"📦 Lote {lote['id']}  ·  {lote.get('data','—')}  ·  "
                    f"{lote.get('total', 0)} cartas  ·  "
                    f"R$ {lote.get('valor_total', 0):,.2f}"
                ):
                    ids_lote = lote.get("ids_cartas", [])
                    cartas_lote = [c for c in cartas if c.get("id") in ids_lote]

                    if cartas_lote:
                        st.dataframe(
                            pd.DataFrame(cartas_lote)[["NOME", "VALOR", "LOJA", "COD_CLI", "DATA"]],
                            use_container_width=True,
                        )

                        col_z, col_x = st.columns(2)

                        zip_hist = gerar_zip_lote(cartas_lote)
                        col_z.download_button(
                            "📥 ZIP do Lote",
                            data=zip_hist,
                            file_name=f"Lote_{lote['id']}.zip",
                            mime="application/zip",
                            key=f"zip_hist_{lote['id']}",
                            use_container_width=True,
                        )

                        xls_hist = gerar_excel_lote(cartas_lote)
                        col_x.download_button(
                            "📊 Excel do Lote",
                            data=xls_hist,
                            file_name=f"Lote_{lote['id']}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"xls_hist_{lote['id']}",
                            use_container_width=True,
                        )
                    else:
                        st.caption("Cartas deste lote não encontradas no banco.")

    # ── ABA 4: CONFIG ────────────────────────────────────────────────────────
    with tabs[4]:
        if user_role not in ["ADM", "GERENTE"]:
            st.warning("🔒 Acesso restrito a Administradores e Gerentes.")
            return

        st.subheader("⚙️ Configurações")

        # ── Cadastro de colaboradores ───────────────────────────────────────
        with st.expander("👤 Gerenciar Colaboradores", expanded=True):
            st.caption(f"{len(dict_colab)} colaborador(es) cadastrado(s)")

            with st.form("form_colab", clear_on_submit=True):
                cc1, cc2 = st.columns(2)
                novo_nome_c = cc1.text_input("Nome *")
                novo_cpf_c  = cc2.text_input("CPF *")
                if st.form_submit_button("➕ Adicionar", use_container_width=True):
                    if novo_nome_c and novo_cpf_c:
                        salvar_novo_colaborador(fire, novo_nome_c, novo_cpf_c)
                        st.success(f"✅ {novo_nome_c.upper()} adicionado!")
                        st.rerun()
                    else:
                        st.error("Preencha nome e CPF.")

            if dict_colab:
                st.write("**Colaboradores cadastrados:**")
                for nome_c, cpf_c in sorted(dict_colab.items()):
                    cc1, cc2 = st.columns([4, 1])
                    cc1.write(f"**{nome_c}** — CPF: {cpf_c}")
                    if cc2.button("🗑️", key=f"del_colab_{nome_c}", use_container_width=True):
                        fire.collection("colaboradores_base").document(nome_c).delete()
                        st.toast(f"{nome_c} removido.")
                        st.rerun()

        # ── Exportar tudo ──────────────────────────────────────────────────
        with st.expander("📤 Exportar Base Completa"):
            todas = [c for c in cartas]
            if todas:
                xls_all = gerar_excel_lote(todas)
                st.download_button(
                    "📊 Exportar TODAS as cartas (.xlsx)",
                    data=xls_all,
                    file_name=f"Base_Completa_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            else:
                st.info("Nenhuma carta no banco ainda.")

        # ── Zona de perigo ─────────────────────────────────────────────────
        if user_role == "ADM":
            with st.expander("🚨 Zona de Perigo"):
                st.warning("Reabrir uma carta muda seu status de volta para 'Aguardando Assinatura'.")
                cartas_fechadas = [c for c in cartas if c.get("status") == "LOTE_FECHADO"]
                if cartas_fechadas:
                    nomes_f = [f"{c['NOME']} — {c['DATA']}" for c in cartas_fechadas]
                    idx_sel = st.selectbox("Selecionar carta para reabrir:", range(len(nomes_f)),
                                           format_func=lambda i: nomes_f[i])
                    if st.button("🔓 Reabrir Carta Selecionada", type="primary"):
                        id_reabrir = cartas_fechadas[idx_sel]["id"]
                        fire.collection("cartas_rh").document(id_reabrir).update(
                            {"status": "Aguardando Assinatura", "id_lote": ""}
                        )
                        st.success("Carta reaberta.")
                        st.rerun()
